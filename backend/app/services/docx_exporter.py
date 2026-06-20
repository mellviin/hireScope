"""
DOCX Export Engine
Generates ATS-optimized resumes in DOCX format
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List
from app.utils.logging import setup_logging

logger = setup_logging(__name__)


class DOCXExporter:
    """Export resume to ATS-optimized DOCX format"""
    
    HEADING_SIZE = 14
    SUBHEADING_SIZE = 12
    BODY_SIZE = 11
    HEADING_COLOR = RGBColor(31, 78, 121)  # Professional blue
    
    def __init__(self):
        """Initialize DOCX exporter"""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles"""
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def export_resume(self, parsed_resume: Dict[str, Any], filename: str = "resume.docx") -> bytes:
        """
        Export resume to DOCX format
        
        Args:
            parsed_resume: Parsed resume data
            filename: Output filename
            
        Returns:
            DOCX file bytes
        """
        try:
            # Header with name and contact info
            self._add_header(parsed_resume)
            
            # Add sections
            if parsed_resume.get('summary'):
                self._add_summary(parsed_resume['summary'])
            
            if parsed_resume.get('experience'):
                self._add_experience(parsed_resume['experience'])
            
            if parsed_resume.get('education'):
                self._add_education(parsed_resume['education'])
            
            if parsed_resume.get('skills'):
                self._add_skills(parsed_resume['skills'])
            
            if parsed_resume.get('projects'):
                self._add_projects(parsed_resume['projects'])
            
            if parsed_resume.get('certifications'):
                self._add_certifications(parsed_resume['certifications'])
            
            # Save to bytes
            import io
            output = io.BytesIO()
            self.doc.save(output)
            output.seek(0)
            
            logger.info(f"Resume exported to DOCX: {filename}")
            return output.getvalue()
        except Exception as e:
            logger.error(f"Error exporting resume: {str(e)}")
            raise
    
    def _add_header(self, resume: Dict[str, Any]):
        """Add resume header with contact info"""
        # Name
        if resume.get('name'):
            name_para = self.doc.add_paragraph(resume['name'])
            name_para.style = 'Heading 1'
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.runs[0]
            name_run.font.size = Pt(self.HEADING_SIZE + 2)
            name_run.font.color.rgb = self.HEADING_COLOR
        
        # Contact info
        contact_info = []
        if resume.get('email'):
            contact_info.append(resume['email'])
        if resume.get('phone'):
            contact_info.append(resume['phone'])
        if resume.get('linkedin'):
            contact_info.append(resume['linkedin'])
        if resume.get('github'):
            contact_info.append(resume['github'])
        
        if contact_info:
            contact_para = self.doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(10)
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_summary(self, summary: str):
        """Add professional summary section"""
        self._add_section_heading('PROFESSIONAL SUMMARY')
        self.doc.add_paragraph(summary)
    
    def _add_experience(self, experience: List[Dict[str, Any]]):
        """Add experience section"""
        if not experience:
            return
        
        self._add_section_heading('PROFESSIONAL EXPERIENCE')
        
        for exp in experience:
            # Job title and company
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(exp.get('title', 'Position'))
            title_run.bold = True
            title_run.font.size = Pt(self.SUBHEADING_SIZE)
            
            # Duration
            if exp.get('duration'):
                duration_run = title_para.add_run(f" | {exp['duration']}")
                duration_run.font.size = Pt(10)
                duration_run.font.italic = True
            
            # Company
            if exp.get('company'):
                company_para = self.doc.add_paragraph(exp['company'], style='List Bullet')
            
            # Description
            if exp.get('description'):
                self.doc.add_paragraph(exp['description'], style='List Bullet')
            
            # Skills
            if exp.get('skills'):
                skills_text = f"Skills: {', '.join(exp['skills'])}"
                skills_para = self.doc.add_paragraph(skills_text, style='List Bullet')
                for run in skills_para.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)
    
    def _add_education(self, education: List[Dict[str, Any]]):
        """Add education section"""
        if not education:
            return
        
        self._add_section_heading('EDUCATION')
        
        for edu in education:
            # Degree and field
            degree_para = self.doc.add_paragraph()
            degree_run = degree_para.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}")
            degree_run.bold = True
            degree_run.font.size = Pt(self.SUBHEADING_SIZE)
            
            # Institution and year
            if edu.get('institution'):
                inst_para = self.doc.add_paragraph(f"{edu['institution']}, {edu.get('graduation_year', '')}")
    
    def _add_skills(self, skills: List[Dict[str, Any]]):
        """Add skills section"""
        if not skills:
            return
        
        self._add_section_heading('TECHNICAL SKILLS')
        
        # Group skills by category
        skills_by_category = {}
        for skill in skills:
            category = skill.get('category', 'General')
            if category not in skills_by_category:
                skills_by_category[category] = []
            skill_name = skill.get('name') if isinstance(skill, dict) else skill
            skills_by_category[category].append(skill_name)
        
        # Add skills
        for category, skill_list in skills_by_category.items():
            skills_para = self.doc.add_paragraph()
            category_run = skills_para.add_run(f"{category}: ")
            category_run.bold = True
            skills_run = skills_para.add_run(', '.join(skill_list[:10]))
    
    def _add_projects(self, projects: List[Dict[str, Any]]):
        """Add projects section"""
        if not projects:
            return
        
        self._add_section_heading('PROJECTS')
        
        for project in projects:
            # Project title
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(project.get('title', 'Project'))
            title_run.bold = True
            title_run.font.size = Pt(self.SUBHEADING_SIZE)
            
            # Description
            if project.get('description'):
                self.doc.add_paragraph(project['description'], style='List Bullet')
            
            # Technologies
            if project.get('technologies'):
                tech_para = self.doc.add_paragraph(
                    f"Technologies: {', '.join(project['technologies'])}",
                    style='List Bullet'
                )
                for run in tech_para.runs:
                    run.font.italic = True
    
    def _add_certifications(self, certifications: List[Dict[str, Any]]):
        """Add certifications section"""
        if not certifications:
            return
        
        self._add_section_heading('CERTIFICATIONS')
        
        for cert in certifications:
            name = cert.get('name', 'Certification')
            issuer = cert.get('issuer', '')
            year = cert.get('date_obtained', '')
            
            cert_text = name
            if issuer:
                cert_text += f" - {issuer}"
            if year:
                cert_text += f" ({year})"
            
            self.doc.add_paragraph(cert_text, style='List Bullet')
    
    def _add_section_heading(self, heading: str):
        """Add section heading"""
        heading_para = self.doc.add_paragraph(heading)
        heading_para.style = 'Heading 2'
        
        # Format heading
        for run in heading_para.runs:
            run.font.size = Pt(self.HEADING_SIZE)
            run.font.bold = True
            run.font.color.rgb = self.HEADING_COLOR
        
        # Add line separator
        self.doc.add_paragraph()
