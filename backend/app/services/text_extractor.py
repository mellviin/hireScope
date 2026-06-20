"""
Resume text extraction from PDF and DOCX files
"""
import io
from typing import Tuple, BinaryIO
import pdfplumber
from docx import Document
from app.utils.logging import setup_logging
from app.utils.exceptions import ProcessingError

logger = setup_logging(__name__)


class ResumeTextExtractor:
    """Extract text from PDF and DOCX files"""
    
    @staticmethod
    def extract_from_pdf(file_content: BinaryIO) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_content: PDF file in bytes or file-like object
            
        Returns:
            Extracted text from PDF
        """
        try:
            with pdfplumber.open(io.BytesIO(file_content.read()) if hasattr(file_content, 'read') else file_content) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
            
            if not text.strip():
                raise ProcessingError("No text could be extracted from PDF")
            
            logger.info(f"Successfully extracted text from PDF ({len(text)} characters)")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ProcessingError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_content: BinaryIO) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_content: DOCX file in bytes or file-like object
            
        Returns:
            Extracted text from DOCX
        """
        try:
            # Read file content
            if hasattr(file_content, 'read'):
                file_bytes = file_content.read()
            else:
                file_bytes = file_content
            
            doc = Document(io.BytesIO(file_bytes))
            
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
            
            if not text.strip():
                raise ProcessingError("No text could be extracted from DOCX")
            
            logger.info(f"Successfully extracted text from DOCX ({len(text)} characters)")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ProcessingError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text(file_content: BinaryIO, file_type: str) -> str:
        """
        Extract text from file based on type
        
        Args:
            file_content: File content
            file_type: Type of file (pdf or docx)
            
        Returns:
            Extracted text
        """
        if file_type.lower() == 'pdf':
            return ResumeTextExtractor.extract_from_pdf(file_content)
        elif file_type.lower() in ['docx', 'doc']:
            return ResumeTextExtractor.extract_from_docx(file_content)
        else:
            raise ProcessingError(f"Unsupported file type: {file_type}")
